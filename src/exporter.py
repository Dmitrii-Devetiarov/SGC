"""Экспорт гипотез в DOCX."""

from docx import Document
from datetime import datetime

def export_to_docx(hypotheses: list, query: str, path: str):
    doc = Document()
    doc.add_heading("SGC — Отчёт о гипотезах", 0)

    if query:
        doc.add_paragraph(f"Запрос: {query}")
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Гипотез: {len(hypotheses)}")

    for h in hypotheses:
        doc.add_heading(f"Гипотеза #{h.get('gap_id', '?')} | Глубина разрыва: {h.get('gap_depth', 0):.3f}", 2)
        doc.add_paragraph(h.get("hypothesis", ""))
        doc.add_heading("Механизм", 3)
        doc.add_paragraph(h.get("mechanism", ""))
        doc.add_heading("Обоснование", 3)
        doc.add_paragraph(h.get("justification", ""))
        doc.add_heading("Новизна", 3)
        doc.add_paragraph(h.get("novelty", ""))
        doc.add_heading("Риски", 3)
        doc.add_paragraph(h.get("risks", ""))
        doc.add_heading("План проверки", 3)
        doc.add_paragraph(h.get("verification_plan", ""))
        if h.get("sources"):
            doc.add_heading("Источники", 3)
            for s in h["sources"]:
                doc.add_paragraph(s, style="List Bullet")

    doc.save(path)